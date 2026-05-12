"""
agent.py  —  ADK Automated Weekly Report Tool
Production-ready multi-agent pipeline: SOW → Excel (Task Tracker only)

Note: Gantt / Timeline sheet is handled from the UI only.
      This pipeline produces Sheet1 (Task Tracker) + Google Sheets task tracker.

Architecture (SequentialAgent):
  1. SOWParserAgent     → ingests PDF / DOCX / pasted text  → session["sow_raw_text"]
  2. PlannerAgent       → extracts project_data, tasks_data
  3. ExcelBuilderAgent  → Google Sheets task tracker with dropdown statuses + conditional colors
  4. ReportSummaryAgent → final markdown summary + Google Sheet URL
"""

import json
import base64
import os
import re
import urllib.parse
from collections import OrderedDict
from datetime import datetime, timedelta
from typing import Optional

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from dotenv import load_dotenv
from google.adk.agents import LlmAgent, SequentialAgent
from google.adk.tools import ToolContext
from google.auth.transport.requests import Request
from google.genai import types as genai_types
from google.oauth2.credentials import Credentials as UserCredentials
from google.oauth2.service_account import Credentials as ServiceAccountCredentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

# ══════════════════════════════════════════════════════════════════════════════
# ENVIRONMENT
# ══════════════════════════════════════════════════════════════════════════════

load_dotenv()

OUTPUT_DIR                      = os.environ.get("OUTPUT_DIR",                      "./outputs")
FILE_SERVER_BASE_URL            = os.environ.get("FILE_SERVER_BASE_URL",            "http://localhost:8000/files")
GEMINI_MODEL                    = os.environ.get("GEMINI_MODEL",                    "gemini-2.5-flash")
RETRY_ATTEMPTS                  = int(os.environ.get("RETRY_ATTEMPTS",              "5"))
RETRY_DELAY                     = float(os.environ.get("RETRY_INITIAL_DELAY",       "2"))
GOOGLE_SHEETS_CREDENTIALS_JSON  = os.environ.get("GOOGLE_SHEETS_CREDENTIALS_JSON",  "").strip()
GOOGLE_SHEETS_CREDENTIALS_B64   = os.environ.get("GOOGLE_SHEETS_CREDENTIALS_B64",   "").strip()
GOOGLE_SHEETS_PARENT_FOLDER_ID  = os.environ.get("GOOGLE_SHEETS_PARENT_FOLDER_ID",  "").strip()
GOOGLE_SHEETS_SHARE_WITH        = os.environ.get("GOOGLE_SHEETS_SHARE_WITH",        "").strip()
GOOGLE_SHEETS_SHARE_ROLE        = os.environ.get("GOOGLE_SHEETS_SHARE_ROLE",        "writer").strip() or "writer"
GOOGLE_SHEETS_MAKE_PUBLIC       = os.environ.get("GOOGLE_SHEETS_MAKE_PUBLIC",       "").strip().lower()
GOOGLE_OAUTH_CLIENT_SECRET_JSON = os.environ.get("GOOGLE_OAUTH_CLIENT_SECRET_JSON", "").strip()
GOOGLE_OAUTH_TOKEN_JSON         = os.environ.get("GOOGLE_OAUTH_TOKEN_JSON",         "./google_oauth_token.json").strip()
GOOGLE_OAUTH_AUTH_MODE          = os.environ.get("GOOGLE_OAUTH_AUTH_MODE",          "local_server").strip().lower()
DEFAULT_SHEETS_CREDENTIALS_JSON = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "sheets_credential.json")
)

os.makedirs(OUTPUT_DIR, exist_ok=True)

_RETRY_CONFIG = genai_types.GenerateContentConfig(
    http_options=genai_types.HttpOptions(
        retry_options=genai_types.HttpRetryOptions(
            initial_delay=RETRY_DELAY,
            attempts=RETRY_ATTEMPTS,
        ),
    ),
)

GOOGLE_SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/spreadsheets",
]

# ══════════════════════════════════════════════════════════════════════════════
# EXCEL STYLING  (Sheet1 — Task Tracker)
# ══════════════════════════════════════════════════════════════════════════════

HEADER_FILL = "FFD9D9D9"
BLACK       = "FF000000"


def _thin_border() -> Border:
    s = Side(style="thin", color=BLACK)
    return Border(left=s, right=s, top=s, bottom=s)


def _hdr(cell, value: str, size: int = 12):
    cell.value     = value
    cell.font      = Font(bold=True, size=size, name="Calibri", color=BLACK)
    cell.fill      = PatternFill("solid", fgColor=HEADER_FILL)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border    = _thin_border()


def _data(cell, value, align_h: str = "left", fmt: str = None):
    cell.value     = value
    cell.font      = Font(bold=False, size=12, name="Calibri", color=BLACK)
    cell.fill      = PatternFill(fill_type=None)
    cell.alignment = Alignment(horizontal=align_h, vertical="center", wrap_text=True)
    cell.border    = _thin_border()
    if fmt:
        cell.number_format = fmt


def _set_col_width(ws, col_letter: str, width: float):
    ws.column_dimensions[col_letter].width = width

# ══════════════════════════════════════════════════════════════════════════════
# UTILITIES
# ══════════════════════════════════════════════════════════════════════════════

def _safe_json(value):
    if isinstance(value, (dict, list)):
        return value
    if not value:
        return None
    try:
        return json.loads(value)
    except (json.JSONDecodeError, TypeError):
        return None


def _parse_date(value) -> Optional[datetime]:
    if not value:
        return None
    if isinstance(value, datetime):
        return value
    raw = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y",
                "%d %b %Y", "%d %B %Y", "%Y/%m/%d", "%b %d, %Y"):
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue
    return None


def _slugify(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", text.strip())[:40]


def _format_task_bullets(text: str) -> str:
    if not text:
        return ""
    parts = [p.strip() for p in text.split(" | ") if p.strip()]
    return ("• " + "\n• ".join(parts)) if len(parts) > 1 else text


def _normalize_status(status: str) -> str:
    value = (status or "").strip().lower()
    if value in {"completed", "complete", "done", "finished"}:
        return "Completed"
    if value in {"todo", "to do", "to-do", "pending"}:
        return "Todo"
    if value in {"in progress", "inprogress", "wip", "work in progress", "ongoing"}:
        return "In Progress"
    if value in {"delayed", "delay", "blocked"}:
        return "Delayed"
    return "Todo"


def _summarize_subtasks(task: dict) -> str:
    raw = str(
        task.get("subtasks_summary")
        or task.get("subtasks")
        or task.get("task_detail")
        or task.get("detail")
        or ""
    ).strip()
    if not raw:
        return ""
    parts = [p.strip("• -\t") for p in re.split(r"\s*\|\s*|\n+", raw) if p.strip()]
    if not parts:
        return ""
    if len(parts) == 1:
        text = parts[0]
        return text if len(text) <= 180 else (text[:177] + "...")
    first  = parts[0]
    second = parts[1]
    summary = f"{first}. {second}" if not first.endswith((".", "!", "?")) else f"{first} {second}"
    return summary if len(summary) <= 220 else (summary[:217] + "...")


def _resolve_google_path(path_value: str) -> str:
    if not path_value:
        return ""
    if os.path.isabs(path_value):
        return path_value
    return os.path.abspath(os.path.join(os.path.dirname(__file__), "..", path_value))


def _load_google_credentials():
    oauth_secret_path = _resolve_google_path(GOOGLE_OAUTH_CLIENT_SECRET_JSON)
    oauth_token_path  = _resolve_google_path(GOOGLE_OAUTH_TOKEN_JSON)

    if oauth_secret_path:
        if not os.path.exists(oauth_secret_path):
            raise FileNotFoundError(
                f"GOOGLE_OAUTH_CLIENT_SECRET_JSON not found: {oauth_secret_path}"
            )
        creds = None
        if oauth_token_path and os.path.exists(oauth_token_path):
            try:
                creds = UserCredentials.from_authorized_user_file(oauth_token_path, GOOGLE_SCOPES)
            except Exception:
                creds = None
        
        needs_reauth = False
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception:
                needs_reauth = True
        elif not creds or not creds.valid:
            needs_reauth = True
        
        if needs_reauth:
            if os.path.exists(oauth_token_path):
                try:
                    os.remove(oauth_token_path)
                except Exception:
                    pass
            flow = InstalledAppFlow.from_client_secrets_file(
                oauth_secret_path, GOOGLE_SCOPES
            )
            if GOOGLE_OAUTH_AUTH_MODE == "console":
                creds = flow.run_console()
            else:
                try:
                    creds = flow.run_local_server(port=0)
                except Exception:
                    creds = flow.run_console()
        
        if oauth_token_path and creds:
            os.makedirs(os.path.dirname(oauth_token_path) or ".", exist_ok=True)
            with open(oauth_token_path, "w", encoding="utf-8") as f:
                f.write(creds.to_json())
        return creds

    json_path = GOOGLE_SHEETS_CREDENTIALS_JSON
    if not json_path and os.path.exists(DEFAULT_SHEETS_CREDENTIALS_JSON):
        json_path = DEFAULT_SHEETS_CREDENTIALS_JSON

    if json_path:
        json_path = _resolve_google_path(json_path)
        if not os.path.exists(json_path):
            raise FileNotFoundError(
                f"GOOGLE_SHEETS_CREDENTIALS_JSON not found: {json_path}"
            )
        with open(json_path, "r", encoding="utf-8") as f:
            info = json.load(f)
        return ServiceAccountCredentials.from_service_account_info(info, scopes=GOOGLE_SCOPES)

    if GOOGLE_SHEETS_CREDENTIALS_B64:
        decoded = base64.b64decode(GOOGLE_SHEETS_CREDENTIALS_B64).decode("utf-8")
        info    = json.loads(decoded)
        return ServiceAccountCredentials.from_service_account_info(info, scopes=GOOGLE_SCOPES)

    raise ValueError(
        "Missing Google credentials. Set GOOGLE_OAUTH_CLIENT_SECRET_JSON for personal "
        "Drive, or GOOGLE_SHEETS_CREDENTIALS_JSON / GOOGLE_SHEETS_CREDENTIALS_B64 "
        "for service account mode."
    )


def _parse_share_emails(raw: str) -> list[str]:
    return [email.strip() for email in raw.split(",") if email.strip()]

# ══════════════════════════════════════════════════════════════════════════════
# TOOLS — SOW INGESTION
# ══════════════════════════════════════════════════════════════════════════════

def read_pdf_sow(file_path: str, *, tool_context: ToolContext) -> dict:
    """
    Extract full text from a SOW PDF file and store in session state.

    Args:
        file_path: Absolute or relative path to a .pdf SOW document.

    Returns:
        success (bool), pages (int), char_count (int), error (str)
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return {"success": False, "pages": 0, "char_count": 0,
                "error": "pypdf not installed — run: pip install pypdf"}
    path = os.path.abspath(os.path.expanduser(file_path.strip()))
    if not os.path.exists(path):
        return {"success": False, "pages": 0, "char_count": 0,
                "error": f"File not found: {path}"}
    if not path.lower().endswith(".pdf"):
        return {"success": False, "pages": 0, "char_count": 0,
                "error": "Not a PDF. For .docx use read_docx_sow."}
    try:
        reader = PdfReader(path)
        text   = "\n\n".join(
            f"--- Page {i+1} ---\n{(p.extract_text() or '').strip()}"
            for i, p in enumerate(reader.pages)
        )
        tool_context.state["sow_raw_text"] = text
        tool_context.state["sow_source"]   = f"PDF:{os.path.basename(path)}"
        return {"success": True, "pages": len(reader.pages),
                "char_count": len(text), "error": ""}
    except Exception as exc:
        return {"success": False, "pages": 0, "char_count": 0, "error": str(exc)}


def read_docx_sow(file_path: str, *, tool_context: ToolContext) -> dict:
    """
    Extract full text from a SOW Word (.docx) file and store in session state.

    Args:
        file_path: Absolute or relative path to a .docx SOW document.

    Returns:
        success (bool), paragraphs (int), char_count (int), error (str)
    """
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "paragraphs": 0, "char_count": 0,
                "error": "python-docx not installed — run: pip install python-docx"}
    path = os.path.abspath(os.path.expanduser(file_path.strip()))
    if not os.path.exists(path):
        return {"success": False, "paragraphs": 0, "char_count": 0,
                "error": f"File not found: {path}"}
    if not path.lower().endswith(".docx"):
        return {"success": False, "paragraphs": 0, "char_count": 0,
                "error": "Not a .docx file. For PDFs use read_pdf_sow."}
    try:
        doc        = Document(path)
        paras      = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_rows = []
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if cells:
                    table_rows.append(cells)
        text = "\n".join(paras)
        if table_rows:
            text += "\n\n--- Tables ---\n" + "\n".join(table_rows)
        tool_context.state["sow_raw_text"] = text
        tool_context.state["sow_source"]   = f"DOCX:{os.path.basename(path)}"
        return {"success": True, "paragraphs": len(paras),
                "char_count": len(text), "error": ""}
    except Exception as exc:
        return {"success": False, "paragraphs": 0, "char_count": 0, "error": str(exc)}


def store_sow_text(sow_text: str, *, tool_context: ToolContext) -> dict:
    """
    Store raw pasted SOW text in session state.

    Args:
        sow_text: The complete raw text of the SOW document.

    Returns:
        success (bool), char_count (int), error (str)
    """
    cleaned = (sow_text or "").strip()
    if not cleaned:
        return {"success": False, "char_count": 0,
                "error": "sow_text is empty — paste the full SOW content."}
    tool_context.state["sow_raw_text"] = cleaned
    tool_context.state["sow_source"]   = "PastedText"
    return {"success": True, "char_count": len(cleaned), "error": ""}


# ══════════════════════════════════════════════════════════════════════════════
# TOOLS — STRUCTURED PLAN STORAGE
# ══════════════════════════════════════════════════════════════════════════════

def save_project_metadata(project_data_json: str, *, tool_context: ToolContext) -> dict:
    """
    Save project metadata extracted from the SOW to session state.

    Args:
        project_data_json: JSON string with keys:
            project_name, client_name, vendor_name, spoc_name, spoc_email,
            start_date (YYYY-MM-DD), end_date (YYYY-MM-DD), total_weeks (int)

    Returns:
        success (bool), project_name (str), error (str)
    """
    try:
        project = json.loads(project_data_json)
    except json.JSONDecodeError as exc:
        return {"success": False, "project_name": "", "error": f"Invalid JSON: {exc}"}

    missing = {"project_name", "client_name", "start_date", "end_date"} - set(project.keys())
    if missing:
        return {"success": False, "project_name": "",
                "error": f"Missing required keys: {missing}"}

    project.setdefault("vendor_name",  "")
    project.setdefault("spoc_name",    "")
    project.setdefault("spoc_email",   "")
    project.setdefault("total_weeks",  3)

    tool_context.state["project_data"] = project
    return {"success": True, "project_name": project["project_name"], "error": ""}


def save_tasks_data(tasks_data_json: str, *, tool_context: ToolContext) -> dict:
    """
    Save the task list (Sheet1 Task Tracker) to session state.
    Auto-consolidates if more than 20 tasks are submitted.

    Args:
        tasks_data_json: JSON array of task objects, each with:
            phase, module, task_detail, assigned_to,
            start_date (YYYY-MM-DD), end_date (YYYY-MM-DD), status, remark

    Returns:
        success (bool), task_count (int), phase_count (int),
        consolidated (bool), error (str)
    """
    MAX_TASKS = 20
    try:
        tasks = json.loads(tasks_data_json)
    except json.JSONDecodeError as exc:
        return {"success": False, "task_count": 0, "phase_count": 0,
                "consolidated": False, "error": f"Invalid JSON: {exc}"}

    if not isinstance(tasks, list) or not tasks:
        return {"success": False, "task_count": 0, "phase_count": 0,
                "consolidated": False, "error": "Must be a non-empty JSON array."}

    required = {"phase", "module", "task_detail", "start_date", "end_date", "status"}
    for i, t in enumerate(tasks):
        missing_t = required - set(t.keys())
        if missing_t:
            return {"success": False, "task_count": 0, "phase_count": 0,
                    "consolidated": False,
                    "error": f"Task[{i}] missing keys: {missing_t}"}
        t.setdefault("assigned_to", "")
        t.setdefault("remark",      "")

    consolidated = False
    if len(tasks) > MAX_TASKS:
        consolidated = True
        groups: OrderedDict = OrderedDict()
        for t in tasks:
            key = (t["phase"], t.get("assigned_to", ""), t["start_date"], t["end_date"])
            if key not in groups:
                groups[key] = dict(t)
            else:
                combined = f"{groups[key]['task_detail']} | {t['task_detail']}"
                groups[key]["task_detail"] = combined[:500]
        tasks = list(groups.values())

        if len(tasks) > MAX_TASKS:
            pg: OrderedDict = OrderedDict()
            for t in tasks:
                p = t["phase"]
                if p not in pg:
                    pg[p] = dict(t)
                    pg[p]["module"] = p.split(":")[-1].strip() if ":" in p else p
                else:
                    combined = f"{pg[p]['task_detail']} | {t['task_detail']}"
                    pg[p]["task_detail"] = combined[:500]
                    if t["end_date"] > pg[p]["end_date"]:
                        pg[p]["end_date"] = t["end_date"]
            tasks = list(pg.values())

    phases = len({t["phase"] for t in tasks})
    tool_context.state["tasks_data"] = tasks
    return {"success": True, "task_count": len(tasks), "phase_count": phases,
            "consolidated": consolidated, "error": ""}


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — GOOGLE SHEETS TASK TRACKER CREATOR
# ══════════════════════════════════════════════════════════════════════════════

def create_google_task_tracker_sheet(*, tool_context: ToolContext) -> dict:
    """
    Create a structured Google Sheets task tracker using Sheets + Drive APIs.

    Reads: project_data, tasks_data from session state.
    Writes: google_sheet_id, google_sheet_url to session state.

    Returns:
        success (bool), google_sheet_id (str), google_sheet_url (str),
        row_count (int), error (str)
    """
    state   = tool_context.state
    project = _safe_json(state.get("project_data"))
    tasks   = _safe_json(state.get("tasks_data"))

    if not isinstance(project, dict) or not project:
        return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                "row_count": 0,
                "error": "project_data missing — run save_project_metadata first."}
    if not isinstance(tasks, list) or not tasks:
        return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                "row_count": 0,
                "error": "tasks_data missing — run save_tasks_data first."}

    try:
        creds = _load_google_credentials()
    except (FileNotFoundError, ValueError, json.JSONDecodeError, base64.binascii.Error) as exc:
        return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                "row_count": 0, "error": str(exc)}

    try:
        sheets = build("sheets", "v4", credentials=creds)
        drive  = build("drive",  "v3", credentials=creds)
    except HttpError as exc:
        return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                "row_count": 0,
                "error": f"Failed to initialize Google APIs: {exc}"}

    service_account_email = getattr(creds, "service_account_email", "service account")
    if GOOGLE_SHEETS_PARENT_FOLDER_ID:
        try:
            drive.files().get(
                fileId=GOOGLE_SHEETS_PARENT_FOLDER_ID,
                fields="id,name,mimeType",
                supportsAllDrives=True,
            ).execute()
        except HttpError as exc:
            return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                    "row_count": 0,
                    "error": (
                        f"GOOGLE_SHEETS_PARENT_FOLDER_ID is not accessible. "
                        f"Share folder with '{service_account_email}' as Editor. "
                        f"Raw error: {exc}"
                    )}

    client    = _slugify(project.get("client_name", "client"))
    ts        = datetime.now().strftime("%Y%m%d_%H%M%S")
    title     = f"{client}_Task_Tracker_{ts}"

    try:
        create_body = {
            "name":     title,
            "mimeType": "application/vnd.google-apps.spreadsheet",
        }
        if GOOGLE_SHEETS_PARENT_FOLDER_ID:
            create_body["parents"] = [GOOGLE_SHEETS_PARENT_FOLDER_ID]

        created = drive.files().create(
            body=create_body,
            fields="id,webViewLink",
            supportsAllDrives=True,
        ).execute()
    except HttpError as exc:
        err = str(exc)
        if "storage quota" in err.lower() or "quotaExceeded" in err:
            return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                    "row_count": 0,
                    "error": (
                        "Google Drive storage quota exceeded for the service account. "
                        "Use a Shared Drive folder and set GOOGLE_SHEETS_PARENT_FOLDER_ID."
                    )}
        return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                "row_count": 0, "error": f"Google Drive create error: {exc}"}

    spreadsheet_id  = created.get("id", "")
    spreadsheet_url = created.get(
        "webViewLink",
        f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}",
    )

    try:
        sheet_meta = sheets.spreadsheets().get(
            spreadsheetId=spreadsheet_id,
            fields="sheets(properties(sheetId,title))",
        ).execute()
    except HttpError as exc:
        return {"success": False, "google_sheet_id": spreadsheet_id,
                "google_sheet_url": spreadsheet_url, "row_count": 0,
                "error": f"Google Sheets metadata error: {exc}"}

    first_sheet_props  = sheet_meta["sheets"][0]["properties"]
    sheet_id           = first_sheet_props["sheetId"]
    source_sheet_title = first_sheet_props.get("title", "Sheet1")

    # ── Build header + data rows ──────────────────────────────────────────
    rows           = [["Phases", "Module", "Tasks", "Start Date", "End Date",
                       "Duration", "Status", "Remarks"]]
    timeline_tasks = []

    for task in tasks:
        phase      = str(task.get("phase", "")).strip()
        module     = str(task.get("module") or task.get("task_name") or "").strip()
        task_text  = _summarize_subtasks(task) or str(
            task.get("task_detail") or task.get("detail") or ""
        ).strip()
        start_dt   = _parse_date(task.get("start_date"))
        end_dt     = _parse_date(task.get("end_date"))
        status     = _normalize_status(str(task.get("status", "Todo")))
        remarks    = str(task.get("remark", "")).strip()
        duration   = str((end_dt - start_dt).days) if start_dt and end_dt else ""

        if not phase and not module and not task_text:
            continue

        rows.append([
            phase,
            module,
            task_text,
            start_dt.strftime("%Y-%m-%d") if start_dt else "",
            end_dt.strftime("%Y-%m-%d")   if end_dt   else "",
            duration,
            status,
            remarks,
        ])
        timeline_tasks.append({
            "phase":     phase,
            "task_text": task_text,
            "start_dt":  start_dt,
            "end_dt":    end_dt,
        })

    if len(rows) == 1:
        return {"success": False, "google_sheet_id": "", "google_sheet_url": "",
                "row_count": 0, "error": "No non-empty task rows to write."}

    # ── Write data ────────────────────────────────────────────────────────
    try:
        sheets.spreadsheets().values().update(
            spreadsheetId=spreadsheet_id,
            range=f"{source_sheet_title}!A1:H",
            valueInputOption="USER_ENTERED",
            body={"values": rows},
        ).execute()

        last_row = len(rows)
        requests = [
            # Rename sheet
            {"updateSheetProperties": {
                "properties": {"sheetId": sheet_id, "title": "Task Tracker"},
                "fields": "title",
            }},
            # Freeze header row
            {"updateSheetProperties": {
                "properties": {"sheetId": sheet_id,
                               "gridProperties": {"frozenRowCount": 1}},
                "fields": "gridProperties.frozenRowCount",
            }},
            # Header row format
            {"repeatCell": {
                "range": {"sheetId": sheet_id,
                          "startRowIndex": 0, "endRowIndex": 1,
                          "startColumnIndex": 0, "endColumnIndex": 8},
                "cell": {"userEnteredFormat": {
                    "backgroundColor": {"red": 0.85, "green": 0.85, "blue": 0.85},
                    "textFormat": {"bold": True},
                    "horizontalAlignment": "CENTER",
                    "verticalAlignment": "MIDDLE",
                    "wrapStrategy": "WRAP",
                }},
                "fields": "userEnteredFormat(backgroundColor,textFormat,"
                          "horizontalAlignment,verticalAlignment,wrapStrategy)",
            }},
            # Date column format
            {"repeatCell": {
                "range": {"sheetId": sheet_id,
                          "startRowIndex": 1, "endRowIndex": last_row,
                          "startColumnIndex": 3, "endColumnIndex": 5},
                "cell": {"userEnteredFormat": {
                    "numberFormat": {"type": "DATE", "pattern": "yyyy-mm-dd"},
                    "horizontalAlignment": "CENTER",
                }},
                "fields": "userEnteredFormat(numberFormat,horizontalAlignment)",
            }},
            # Status dropdown validation
            {"setDataValidation": {
                "range": {"sheetId": sheet_id,
                          "startRowIndex": 1, "endRowIndex": last_row,
                          "startColumnIndex": 6, "endColumnIndex": 7},
                "rule": {
                    "condition": {
                        "type": "ONE_OF_LIST",
                        "values": [
                            {"userEnteredValue": "Completed"},
                            {"userEnteredValue": "Todo"},
                            {"userEnteredValue": "In Progress"},
                            {"userEnteredValue": "Delayed"},
                        ],
                    },
                    "strict": True, "showCustomUi": True,
                },
            }},
            # Column widths
            {"updateDimensionProperties": {
                "range": {"sheetId": sheet_id, "dimension": "COLUMNS",
                          "startIndex": 0, "endIndex": 1},
                "properties": {"pixelSize": 220}, "fields": "pixelSize",
            }},
            {"updateDimensionProperties": {
                "range": {"sheetId": sheet_id, "dimension": "COLUMNS",
                          "startIndex": 1, "endIndex": 2},
                "properties": {"pixelSize": 220}, "fields": "pixelSize",
            }},
            {"updateDimensionProperties": {
                "range": {"sheetId": sheet_id, "dimension": "COLUMNS",
                          "startIndex": 2, "endIndex": 3},
                "properties": {"pixelSize": 480}, "fields": "pixelSize",
            }},
            {"updateDimensionProperties": {
                "range": {"sheetId": sheet_id, "dimension": "COLUMNS",
                          "startIndex": 3, "endIndex": 6},
                "properties": {"pixelSize": 120}, "fields": "pixelSize",
            }},
            {"updateDimensionProperties": {
                "range": {"sheetId": sheet_id, "dimension": "COLUMNS",
                          "startIndex": 6, "endIndex": 7},
                "properties": {"pixelSize": 140}, "fields": "pixelSize",
            }},
            {"updateDimensionProperties": {
                "range": {"sheetId": sheet_id, "dimension": "COLUMNS",
                          "startIndex": 7, "endIndex": 8},
                "properties": {"pixelSize": 240}, "fields": "pixelSize",
            }},
            # Conditional formatting — status colours
            {"addConditionalFormatRule": {"rule": {
                "ranges": [{"sheetId": sheet_id,
                            "startRowIndex": 1, "endRowIndex": last_row,
                            "startColumnIndex": 6, "endColumnIndex": 7}],
                "booleanRule": {
                    "condition": {"type": "TEXT_EQ",
                                  "values": [{"userEnteredValue": "Completed"}]},
                    "format": {"backgroundColor": {"red": 0.78, "green": 0.94, "blue": 0.81}},
                },
            }, "index": 0}},
            {"addConditionalFormatRule": {"rule": {
                "ranges": [{"sheetId": sheet_id,
                            "startRowIndex": 1, "endRowIndex": last_row,
                            "startColumnIndex": 6, "endColumnIndex": 7}],
                "booleanRule": {
                    "condition": {"type": "TEXT_EQ",
                                  "values": [{"userEnteredValue": "Todo"}]},
                    "format": {"backgroundColor": {"red": 0.87, "green": 0.93, "blue": 0.98}},
                },
            }, "index": 0}},
            {"addConditionalFormatRule": {"rule": {
                "ranges": [{"sheetId": sheet_id,
                            "startRowIndex": 1, "endRowIndex": last_row,
                            "startColumnIndex": 6, "endColumnIndex": 7}],
                "booleanRule": {
                    "condition": {"type": "TEXT_EQ",
                                  "values": [{"userEnteredValue": "In Progress"}]},
                    "format": {"backgroundColor": {"red": 1.0, "green": 0.95, "blue": 0.8}},
                },
            }, "index": 0}},
            {"addConditionalFormatRule": {"rule": {
                "ranges": [{"sheetId": sheet_id,
                            "startRowIndex": 1, "endRowIndex": last_row,
                            "startColumnIndex": 6, "endColumnIndex": 7}],
                "booleanRule": {
                    "condition": {"type": "TEXT_EQ",
                                  "values": [{"userEnteredValue": "Delayed"}]},
                    "format": {"backgroundColor": {"red": 0.96, "green": 0.8, "blue": 0.8}},
                },
            }, "index": 0}},
        ]

        sheets.spreadsheets().batchUpdate(
            spreadsheetId=spreadsheet_id,
            body={"requests": requests},
        ).execute()

    except HttpError as exc:
        return {"success": False, "google_sheet_id": spreadsheet_id,
                "google_sheet_url": spreadsheet_url, "row_count": len(rows) - 1,
                "error": f"Google Sheets update error: {exc}"}

    # ── Share permissions ─────────────────────────────────────────────────
    try:
        for email in _parse_share_emails(GOOGLE_SHEETS_SHARE_WITH):
            drive.permissions().create(
                fileId=spreadsheet_id,
                body={"type": "user", "role": GOOGLE_SHEETS_SHARE_ROLE,
                      "emailAddress": email},
                sendNotificationEmail=False,
                supportsAllDrives=True,
            ).execute()

        if GOOGLE_SHEETS_MAKE_PUBLIC in {"1", "true", "yes"}:
            drive.permissions().create(
                fileId=spreadsheet_id,
                body={"type": "anyone", "role": "reader"},
                sendNotificationEmail=False,
                supportsAllDrives=True,
            ).execute()
    except HttpError as exc:
        return {"success": False, "google_sheet_id": spreadsheet_id,
                "google_sheet_url": spreadsheet_url, "row_count": len(rows) - 1,
                "error": f"Google Drive sharing error: {exc}"}

    state["google_sheet_id"]  = spreadsheet_id
    state["google_sheet_url"] = spreadsheet_url

    return {"success": True, "google_sheet_id": spreadsheet_id,
            "google_sheet_url": spreadsheet_url,
            "row_count": len(rows) - 1, "error": ""}


# ══════════════════════════════════════════════════════════════════════════════
# TOOL — STATE READER
# ══════════════════════════════════════════════════════════════════════════════

def get_report_state(*, tool_context: ToolContext) -> dict:
    """
    Return a snapshot of all report-related session state for the summary agent.

    Returns:
        project_data (dict), tasks_data (list), excel_file_name (str),
        download_url (str), local_path (str), google_sheet_id (str),
        google_sheet_url (str), success (bool)
    """
    state = tool_context.state
    return {
        "success":          True,
        "project_data":     _safe_json(state.get("project_data"))  or {},
        "tasks_data":       _safe_json(state.get("tasks_data"))    or [],
        "excel_file_name":  state.get("excel_file_name",  ""),
        "download_url":     state.get("download_url",     ""),
        "local_path":       state.get("excel_file_path",  ""),
        "google_sheet_id":  state.get("google_sheet_id",  ""),
        "google_sheet_url": state.get("google_sheet_url", ""),
    }


# ══════════════════════════════════════════════════════════════════════════════
# AGENTS
# ══════════════════════════════════════════════════════════════════════════════

sow_parser_agent = LlmAgent(
    name="SOWParserAgent",
    model=GEMINI_MODEL,
    description=(
        "Ingests a SOW from a PDF path, DOCX path, or raw pasted text "
        "and stores the extracted text in session state."
    ),
    instruction="""
You ingest SOW documents. Make exactly ONE tool call.

Decision rules:
- Message has a path ending in .pdf  → read_pdf_sow(file_path="...")
- Message has a path ending in .docx → read_docx_sow(file_path="...")
- Message has actual SOW text        → store_sow_text(sow_text="...")

Reply with plain text: success/failure, source type, chars/pages extracted, confirmation.
Do not quote SOW content. Do not ask questions.
""",
    tools=[read_pdf_sow, read_docx_sow, store_sow_text],
    generate_content_config=_RETRY_CONFIG,
)


planner_agent = LlmAgent(
    name="PlannerAgent",
    model=GEMINI_MODEL,
    description=(
        "Reads SOW text from session state and extracts project metadata "
        "and task list. Saves via 2 separate tool calls."
    ),
    instruction="""
Extract a structured project plan from session["sow_raw_text"].
Make exactly 2 tool calls in this order.

CALL 1 — save_project_metadata
JSON string keys: project_name, client_name, vendor_name, spoc_name, spoc_email,
start_date (YYYY-MM-DD), end_date (YYYY-MM-DD), total_weeks (int).
If dates are missing use today as start and add total_weeks*7 days for end.

CALL 2 — save_tasks_data
Extract from the Scope of Work section ONLY.
Target: AT MOST 20 task rows total.
Consolidation: group bullets by role within each Track (2-4 rows per Track).
If still over 20, merge entire Track to 1 row with bullets joined by " | ".
Each object: phase, module, task_detail, assigned_to, start_date, end_date, status, remark.
start_date and end_date are actual calendar dates for that track's active window.
status must be one of: "To Do", "In Progress", "Completed", "Delayed".

After both calls succeed reply:
"Plan saved: X tasks across Y tracks. Ready for Google Sheets creation."
If any call fails, report the exact error and stop.
""",
    tools=[save_project_metadata, save_tasks_data],
    generate_content_config=_RETRY_CONFIG,
)


excel_builder_agent = LlmAgent(
    name="ExcelBuilderAgent",
    model=GEMINI_MODEL,
    description=(
        "Creates a Google Sheets task tracker with dropdown statuses and "
        "conditional colours using the Google Sheets and Drive APIs."
    ),
    instruction="""
Create the task tracker Google Sheet from session state. Make exactly 1 tool call:
  create_google_task_tracker_sheet()  — no arguments required

If the tool returns success=False, report the exact error message and stop.
If it succeeds reply:
"Google task tracker created: <google_sheet_url> (<row_count> rows)"
""",
    tools=[create_google_task_tracker_sheet],
    generate_content_config=_RETRY_CONFIG,
)


report_summary_agent = LlmAgent(
    name="ReportSummaryAgent",
    model=GEMINI_MODEL,
    description="Reads final state and produces a clean markdown summary.",
    instruction="""
Call get_report_state() once (no arguments).
Using ONLY values returned by the tool write this markdown summary:

## Weekly Report Generated

**Project:** <project_name>
**Client:** <client_name>  |  **Vendor:** <vendor_name>
**Timeline:** <start_date> to <end_date> (<total_weeks> weeks)

### Contents
| Sheet | Contents |
|-------|----------|
| Task Tracker | <task_count> tasks across N tracks with status dropdowns and conditional colours |

### Track Breakdown
- **<phase>** — <n> tasks  (one bullet per unique phase in tasks_data)

### Output
**Google Sheet:** <google_sheet_url>

Never invent values. Use only what get_report_state() returns.
""",
    tools=[get_report_state],
    generate_content_config=_RETRY_CONFIG,
)


# ══════════════════════════════════════════════════════════════════════════════
# ROOT AGENT
# ══════════════════════════════════════════════════════════════════════════════

root_agent = SequentialAgent(
    name="SOWReportOrchestrator",
    description=(
        "End-to-end pipeline: ingests any SOW (PDF, DOCX, or pasted text), "
        "extracts structured plan, builds a Google Sheets task tracker with "
        "status dropdowns and conditional colours, returns the sheet URL."
    ),
    sub_agents=[
        sow_parser_agent,
        planner_agent,
        excel_builder_agent,
        report_summary_agent,
    ],
)

__all__ = ["root_agent"]


# ══════════════════════════════════════════════════════════════════════════════
# STANDALONE TEST  (python agent.py)
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import asyncio
    from google.adk.runners import Runner
    from google.adk.sessions import InMemorySessionService

    SAMPLE_SOW = """
SOW: KreditBee AWS to GCP Migration POC
Client: KreditBee (Krazybee Services Limited)
Vendor: Wohlig Transformations Pvt Ltd
Start Date: 13 April 2026
End Date: 30 April 2026
Total Timeline: 3 Weeks

Scope of Work

Track 1: Core Integration
- Migrate 18 AWS Lambda functions to Cloud Run across 11 Go binaries
- Configure 50 Pub/Sub topics with push subscriptions
- Provision Apigee org and deploy 7 API proxies

Track 2: Testing & Stability
- Implement GitHub Actions CI with parallel Docker builds
- UI-based pipeline testing, OCR verification, and chunking strategy

Track 3: Infrastructure Scaling
- End-to-End Flow Validation
- GCP Service Bypasses & Integration
- System Stability & Bug Fixes
- Infrastructure Configuration

Track 4: System Hardening
- Customer Demonstration
- Architectural Handover
"""

    async def run():
        svc     = InMemorySessionService()
        session = await svc.create_session(app_name="test", user_id="u1")
        runner  = Runner(agent=root_agent, app_name="test", session_service=svc)
        msg     = genai_types.Content(
            role="user",
            parts=[genai_types.Part(
                text=f"Generate a weekly report from this SOW:\n\n{SAMPLE_SOW}"
            )]
        )
        print("\n" + "=" * 70)
        print(f"  SOW Report Agent  |  Model: {GEMINI_MODEL}")
        print("=" * 70 + "\n")
        async for event in runner.run_async(
                user_id="u1", session_id=session.id, new_message=msg):
            author = getattr(event, "author", "System")
            if event.is_final_response():
                if event.content and event.content.parts:
                    for part in event.content.parts:
                        if getattr(part, "text", ""):
                            print(f"\n{'─'*60}\n  FINAL [{author}]\n{'─'*60}")
                            print(part.text)
            else:
                if event.content and event.content.parts:
                    for part in event.content.parts:
                        txt = (getattr(part, "text", None) or "").strip()
                        if txt:
                            print(f"[{author}] {txt[:140]}")

    asyncio.run(run())